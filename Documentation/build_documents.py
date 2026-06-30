from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageColor, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Final"
ASSETS = ROOT / "Assets"
OUTPUT.mkdir(parents=True, exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)

NAVY = "172B4D"
BLUE = "275DDB"
LIGHT_BLUE = "EAF1FF"
ORANGE = "C2410C"
LIGHT_ORANGE = "FFF3E8"
GREEN = "28734A"
LIGHT_GREEN = "EDF8F1"
INK = "172033"
MUTED = "59677C"
GRID = "DDE3EC"
LIGHT_GRAY = "F4F6FA"
WHITE = "FFFFFF"
BLACK = "000000"

AUTHOR = "Adam Kubiś"
LECTURER = "dr inż. Piotr Górniak"
COURSE = "Kompatybilność elektromagnetyczna"
PROJECT = "EMC Lab Assistant"
VERSION = "2.0"
DATE = "czerwiec 2026"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name="Aptos", size=10.5, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_shading(paragraph, fill: str, border: str | None = None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        for edge in ("top", "left", "bottom", "right"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "6")
            node.set(qn("w:space"), "4")
            node.set(qn("w:color"), border)
            p_bdr.append(node)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Strona ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_document(doc: Document, short_title: str, preset="standard"):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5 if preset == "standard" else 10.25)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5 if preset == "standard" else 4)
    normal.paragraph_format.line_spacing = 1.12 if preset == "standard" else 1.18

    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 10),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 16, NAVY, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(f"{PROJECT}  |  {short_title}")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer.paragraphs[0].clear()
    add_page_number(footer.paragraphs[0])

    doc.core_properties.title = short_title
    doc.core_properties.subject = PROJECT
    doc.core_properties.author = AUTHOR
    doc.core_properties.keywords = "C#, Avalonia UI, EMC, kompatybilność elektromagnetyczna"


def add_cover(doc: Document, title: str, subtitle: str, doc_type: str, metadata: Sequence[tuple[str, str]]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KOMPATYBILNOŚĆ ELEKTROMAGNETYCZNA")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    p = doc.add_paragraph()
    p.style = doc.styles["Subtitle"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(doc_type.upper())
    set_run_font(r, size=10, color=ORANGE, bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(32)

    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(1.05)
        p.paragraph_format.right_indent = Inches(0.75)
        p.paragraph_format.space_after = Pt(3)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=9.8, color=NAVY, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=9.8, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Politechnika Poznańska  •  " + DATE)
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, bold_prefix: str | None = None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullets(doc: Document, items: Iterable[str], compact=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.36)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(2 if compact else 4)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        set_run_font(r)


def add_steps(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.42)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(item)
        set_run_font(r)


def add_callout(doc: Document, label: str, text: str, kind="info"):
    colors = {
        "info": (LIGHT_BLUE, BLUE),
        "warning": (LIGHT_ORANGE, ORANGE),
        "success": (LIGHT_GREEN, GREEN),
    }
    fill, border = colors[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, fill, border)
    r = p.add_run(label + ": ")
    set_run_font(r, color=border, bold=True)
    r = p.add_run(text)
    set_run_font(r, color=INK)


def add_equation(doc: Document, equation: str, explanation: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    set_paragraph_shading(p, LIGHT_GRAY, GRID)
    r = p.add_run(equation)
    set_run_font(r, name="Cambria Math", size=11.5, color=NAVY, bold=True)
    if explanation:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(7)
        r2 = p2.add_run(explanation)
        set_run_font(r2, size=9, color=MUTED, italic=True)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths_dxa: Sequence[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, NAVY)
        cell.text = header
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=9, color=WHITE, bold=True)

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            if row_index % 2:
                shade_cell(cells[index], "F8FAFD")
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, size=8.8)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_image(doc: Document, image_path: Path, caption: str, width=6.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    doc_pr = doc.inline_shapes[-1]._inline.docPr
    doc_pr.set("descr", caption)
    doc_pr.set("title", image_path.stem.replace("_", " "))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption)
    set_run_font(r, size=8.8, color=MUTED, italic=True)


def add_code(doc: Document, lines: Sequence[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.14)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, "F2F4F7", GRID)
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=8.7, color=INK)


def load_font(size: int, bold=False):
    choices = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/seguisb.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
    ]
    for path in choices:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_centered_text(draw, rect, text, font, fill, max_width=None):
    x1, y1, x2, y2 = rect
    words = text.split()
    lines, line = [], ""
    target = max_width or (x2 - x1 - 30)
    for word in words:
        candidate = (line + " " + word).strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= target:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    heights = [draw.textbbox((0, 0), item, font=font)[3] for item in lines]
    total = sum(heights) + max(0, len(lines) - 1) * 5
    y = y1 + (y2 - y1 - total) / 2
    for item, height in zip(lines, heights):
        bbox = draw.textbbox((0, 0), item, font=font)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), item, font=font, fill=fill)
        y += height + 5


def create_diagrams():
    for color in [
        NAVY, BLUE, LIGHT_BLUE, ORANGE, LIGHT_ORANGE, GREEN, LIGHT_GREEN,
        INK, MUTED, GRID, LIGHT_GRAY, WHITE, BLACK, "F8FAFD", "FAFBFD",
        "8FB3FF", "D9E4F7", "AFC0DA", "30486F", "4F8CFF", "E9EEF8",
        "F3E8FF", "6D28D9",
    ]:
        ImageColor.colormap[color.lower()] = f"#{color}"

    font = load_font(28)
    small = load_font(22)
    bold = load_font(28, True)
    title = load_font(34, True)

    # Architecture
    img = Image.new("RGB", (1600, 850), WHITE)
    d = ImageDraw.Draw(img)
    d.text((60, 40), "Architektura logiczna EMC Lab Assistant", font=title, fill=NAVY)
    blocks = [
        ((80, 150, 1520, 280), "Warstwa prezentacji (Avalonia XAML)\nMainWindow, wybór scenariusza, widoki kroków", LIGHT_BLUE, BLUE),
        ((80, 340, 1520, 490), "Warstwa ViewModel (MVVM)\nnawigacja, walidacja, stan kreatorów, komendy", "F4F6FA", NAVY),
        ((80, 550, 740, 720), "Modele danych\npunkty pomiarowe, wyniki, statystyki", LIGHT_GREEN, GREEN),
        ((860, 550, 1520, 720), "Usługi domenowe\nobliczenia, import CSV/MAT, eksport CSV/DOCX", LIGHT_ORANGE, ORANGE),
    ]
    for rect, text, fill, outline in blocks:
        d.rounded_rectangle(rect, radius=22, fill=fill, outline=outline, width=4)
        draw_centered_text(d, rect, text, bold if rect[1] < 500 else font, NAVY)
    for y1, y2 in [(280, 340), (490, 550)]:
        d.line((800, y1 + 10, 800, y2 - 12), fill=MUTED, width=5)
        d.polygon([(790, y2 - 22), (810, y2 - 22), (800, y2 - 8)], fill=MUTED)
    img.save(ASSETS / "architektura.png")

    # Scenario flow
    img = Image.new("RGB", (1600, 1080), WHITE)
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Przebieg czterech kreatorów pomiarowych", font=title, fill=NAVY)
    start = (560, 110, 1040, 205)
    d.rounded_rectangle(start, radius=28, fill=NAVY)
    draw_centered_text(d, start, "Wybór kreatora", bold, WHITE)
    scenario_cards = [
        (80, 285, "Przeniki mikropaskowe", LIGHT_BLUE, BLUE,
         ["1. Dane NEXT/FEXT", "2. Skala liniowa", "3. Błąd analizatora", "4. Statystyka i wykres"]),
        (460, 285, "Sondy pola bliskiego", LIGHT_ORANGE, ORANGE,
         ["1. Stanowisko", "2. Pomiary 30/50/100 Ω", "3. Pole H i U95", "4. Porównanie trendów"]),
        (840, 285, "Emisja EN55032", LIGHT_GREEN, GREEN,
         ["1. Założenia", "2. MR, IL i wysokości", "3. AF i pole E", "4. Limit i margines"]),
        (1220, 285, "Pomiary propagacyjne", "F3E8FF", "6D28D9",
         ["1. Stanowisko DVB-T", "2. Siatka 16 punktów", "3. Pole E", "4. Eav ± T"]),
    ]
    for x, y, heading, fill, outline, steps in scenario_cards:
        rect = (x, y, x + 300, y + 78)
        d.rounded_rectangle(rect, radius=20, fill=fill, outline=outline, width=4)
        draw_centered_text(d, rect, heading, small, NAVY, max_width=260)
        for i, step in enumerate(steps):
            sy = 410 + i * 112
            step_rect = (x, sy, x + 300, sy + 72)
            d.rounded_rectangle(step_rect, radius=16, fill=fill, outline=outline, width=2)
            draw_centered_text(d, step_rect, step, small, INK, max_width=255)
        d.line((800, 205, x + 150, 285), fill=MUTED, width=4)
    export_rect = (560, 965, 1040, 1040)
    d.rounded_rectangle(export_rect, radius=20, fill=LIGHT_GREEN, outline=GREEN, width=4)
    draw_centered_text(d, export_rect, "Eksport wyników do CSV lub DOCX", bold, GREEN)
    for x, *_ in scenario_cards:
        d.line((x + 150, 930, 800, 965), fill=MUTED, width=4)
    img.save(ASSETS / "przeplyw_scenariuszy.png")

    # UI map
    img = Image.new("RGB", (1600, 930), "F4F6FA")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((45, 45, 1555, 885), radius=24, fill=WHITE, outline=GRID, width=4)
    d.rectangle((45, 45, 1555, 220), fill=NAVY)
    d.text((95, 77), "EMC LAB ASSISTANT", font=small, fill="8FB3FF")
    d.text((95, 120), "Tytuł bieżącego kroku", font=title, fill=WHITE)
    d.text((95, 170), "Nazwa scenariusza i krótka instrukcja", font=small, fill="D9E4F7")
    d.rounded_rectangle((1160, 100, 1480, 145), radius=12, fill="30486F")
    d.rounded_rectangle((1160, 100, 1320, 145), radius=12, fill="4F8CFF")
    d.text((1210, 158), "Pasek postępu", font=small, fill=WHITE)
    d.rounded_rectangle((95, 260, 1505, 735), radius=18, fill="FAFBFD", outline=GRID, width=3)
    d.text((125, 290), "Obszar roboczy", font=bold, fill=NAVY)
    for i, label in enumerate(["tabela danych", "wzór / informacja", "wykres / podsumowanie"]):
        x = 145 + i * 440
        d.rounded_rectangle((x, 365, x + 370, 640), radius=16, fill=LIGHT_BLUE if i == 0 else LIGHT_GRAY, outline=GRID, width=2)
        draw_centered_text(d, (x, 365, x + 370, 640), label, font, INK)
    d.rectangle((45, 770, 1555, 885), fill=WHITE)
    d.rounded_rectangle((90, 800, 255, 855), radius=12, fill="E9EEF8")
    d.text((125, 812), "Wstecz", font=small, fill=NAVY)
    d.rounded_rectangle((1305, 800, 1490, 855), radius=12, fill=BLUE)
    d.text((1360, 812), "Dalej", font=small, fill=WHITE)
    img.save(ASSETS / "mapa_interfejsu.png")


def technical_document():
    doc = Document()
    setup_document(doc, "Dokumentacja techniczna", "standard")
    add_cover(
        doc,
        "Dokumentacja techniczna",
        "Architektura, algorytmy, budowanie i utrzymanie programu",
        "Dokumentacja oprogramowania",
        [
            ("Projekt", PROJECT),
            ("Autor", AUTHOR),
            ("Wersja", VERSION),
            ("Technologia", "C# 12, .NET 8, Avalonia UI 12, MVVM"),
            ("Platformy", "Windows 10/11 oraz Linux x64"),
            ("Stan", "Wersja działająca, zweryfikowana kompilacją i testami"),
        ],
    )

    add_heading(doc, "1. Cel i zakres dokumentu")
    add_body(doc, "Dokument opisuje budowę programu EMC Lab Assistant, jego odpowiedzialności funkcjonalne, architekturę, modele danych, algorytmy obliczeniowe, sposób budowania, testowania i publikowania. Jest przeznaczony dla osoby rozwijającej lub oceniającej kod źródłowy.")
    add_callout(doc, "Zakres wersji 2.0", "Program obsługuje cztery operacyjne scenariusze laboratoryjne, moduł nauki oraz ekran pokrycia materiału. Dwa ćwiczenia wymienione w sylabusie oczekują na instrukcje źródłowe i nie są prezentowane jako gotowe.", "info")

    add_heading(doc, "2. Wymagania systemowe")
    add_table(
        doc,
        ["Obszar", "Wymaganie", "Realizacja"],
        [
            ("Język", "C# i środowisko .NET", "Projekt SDK-style, TargetFramework net8.0"),
            ("Interfejs", "Graficzny kreator krok po kroku", "Avalonia XAML, cztery kroki w każdym scenariuszu"),
            ("Wieloplatformowość", "Windows i Linux", "Avalonia Desktop oraz publikacja self-contained"),
            ("Walidacja", "Blokada przejścia przy niepełnych danych", "Warunki CanGoNext i komendy MVVM"),
            ("Obliczenia", "Wzory laboratoryjne i niepewność", "Oddzielne usługi obliczeniowe dla czterech scenariuszy"),
            ("Prezentacja", "Tabele, wzory i wykresy", "DataGrid, CSharpMath oraz własne kontrolki wykresów"),
            ("Eksport", "Dane do sprawozdania", "CSV UTF-8 oraz sformatowany raport DOCX"),
        ],
        [1600, 3300, 4460],
    )

    add_heading(doc, "3. Stos technologiczny")
    add_table(
        doc,
        ["Składnik", "Wersja", "Zastosowanie"],
        [
            ("Microsoft .NET", "8.0", "Runtime, biblioteka standardowa i narzędzia publikowania"),
            ("Avalonia", "12.0.5", "Wieloplatformowy interfejs użytkownika"),
            ("Avalonia.Controls.DataGrid", "12.0.1", "Edycja i prezentacja tabel pomiarowych"),
            ("CommunityToolkit.Mvvm", "8.4.1", "ObservableObject, komendy i generatory właściwości"),
            ("CSharpMath.Avalonia", "12.0.0", "Skład równań matematycznych w interfejsie"),
            ("DocumentFormat.OpenXml", "3.5.1", "Generowanie raportów DOCX"),
            ("MatFileHandler", "1.3.0", "Import wyników MATLAB ćwiczenia nr 3"),
            ("Inter", "pakiet Avalonia", "Spójna typografia interfejsu"),
        ],
        [2400, 1200, 5760],
    )
    add_body(doc, "Avalonia została wybrana zamiast WPF i Windows Forms, ponieważ nie wiąże warstwy prezentacji wyłącznie z systemem Windows. Projekt celuje w .NET 8, dlatego oficjalnym zakresem zgodności są Windows 10/11 oraz wspierane dystrybucje Linuksa.")

    add_heading(doc, "4. Architektura")
    add_image(doc, ASSETS / "architektura.png", "Rysunek 1. Warstwy logiczne programu.")
    add_body(doc, "Aplikacja stosuje wzorzec MVVM. Widoki XAML odpowiadają za prezentację, ViewModele przechowują stan i nawigację, modele opisują rekordy pomiarowe, a usługi wykonują obliczenia oraz eksport. ViewLocator dobiera widok na podstawie typu ViewModelu.")
    add_table(
        doc,
        ["Warstwa / moduł", "Odpowiedzialność"],
        [
            ("MainWindowViewModel", "Wybór scenariusza, bieżący krok, komendy Wstecz/Dalej, restart i zmiana scenariusza."),
            ("ScenarioSelectionViewModel", "Udostępnienie czterech kreatorów, modułu nauki i rejestru pokrycia."),
            ("Step1–Step4ViewModel", "Stan i obliczenia kreatora pomiaru przeników."),
            ("NearFieldStep1–Step4ViewModel", "Stanowisko, dane pomiarowe, pole H, niepewność i podsumowanie sond."),
            ("RadiatedEmissionStep1–Step4ViewModel", "Założenia EN55032, surowe MR/IL, poprawka antenowa, limit i margines."),
            ("PropagationStep1–Step4ViewModel", "Stanowisko DVB-T, siatka 16 punktów, pole E, średnia przestrzenna i tolerancja."),
            ("CrosstalkLogic", "Konwersja dB, błąd analizatora, statystyka i przedział ufności."),
            ("NearFieldLogic", "Przeliczenie mocy na H, niepewność złożona, maksimum i regresja trendu."),
            ("RadiatedEmissionLogic", "Antenna Factor dipola, korekta pionowa, pole E, U95 i margines EN55032."),
            ("PropagationLogic", "Konwersja odczytów na µV, pole E, średnia przestrzenna i tolerancja Eav ± T."),
            ("LearningViewModel", "Siedem bloków wykładowych, pytania kontrolne i trzy kalkulatory EMC."),
            ("SourceRequirementsViewModel", "Jawne wskazanie materiałów, których nie wolno odtwarzać bez instrukcji."),
            ("MeasurementImportService", "Walidowany import CSV/TXT dla czterech scenariuszy."),
            ("MatRadiatedEmissionImporter", "Import struktury Data z plików MATLAB ćwiczenia nr 3."),
            ("ReportGenerator / DocxReportGenerator", "Eksport wyników do CSV oraz raportu DOCX."),
            ("Kontrolki wykresów i wzorów", "Wykresy, mapy cieplne oraz skład LaTeX przez CSharpMath."),
        ],
        [2750, 6610],
    )

    add_heading(doc, "5. Nawigacja i cykl danych")
    add_image(doc, ASSETS / "przeplyw_scenariuszy.png", "Rysunek 2. Nawigacja od wyboru scenariusza do eksportu.")
    add_steps(doc, [
        "Użytkownik wybiera scenariusz na ekranie startowym.",
        "MainWindowViewModel ustawia właściwy ViewModel pierwszego kroku.",
        "Dane wejściowe są walidowane na bieżąco; komenda Dalej jest aktywna dopiero po spełnieniu wymagań.",
        "Przy przejściu do następnego kroku tworzona jest kolekcja wyników obliczeniowych.",
        "Krok końcowy agreguje statystyki, maksima i trendy oraz udostępnia eksport CSV i DOCX.",
    ])

    add_heading(doc, "6. Modele danych")
    add_table(
        doc,
        ["Typ", "Najważniejsze dane"],
        [
            ("BandDefinition", "Nazwa pasma, częstotliwość początkowa i końcowa, niepewność analizatora."),
            ("MeasurementPointViewModel", "Częstotliwość oraz surowe wartości NEXT/FEXT."),
            ("MeasurementResult", "Wartości dB, liniowe, Delta Z i granice wyniku."),
            ("StatisticsResult", "Średnia, s, błąd standardowy i 95% przedział ufności."),
            ("NearFieldMeasurementPoint", "f, moce dla 30/50/100 Ω, K oraz Sp."),
            ("NearFieldResult", "Pole H w dBA/m i A/m oraz U95 dla trzech impedancji."),
            ("NearFieldSummary", "Maksimum, częstotliwość maksimum i nachylenie charakterystyki."),
            ("RadiatedEmissionMeasurementPoint", "f, IL, MR i wysokości anteny dla polaryzacji H/V."),
            ("RadiatedEmissionResult", "AF, korekta pionowa, pola E, limit EN55032, przedział i margines."),
            ("RadiatedEmissionSummary", "Liczba punktów, liczba przekroczeń i punkt krytyczny."),
            ("PropagationMeasurementPoint", "Numer punktu i poziomy dla polaryzacji poziomej oraz pionowej."),
            ("PropagationResult", "U[µV], E[µV/m], E[dBµV/m] i silniejsza polaryzacja."),
            ("PropagationSummary", "Eav, uE, tolerancja, przedział i punkt maksimum."),
        ],
        [2700, 6660],
    )

    add_heading(doc, "7. Algorytmy scenariusza przeników")
    add_equation(doc, "|Z|lin = 10^(|Z|dB / 20)", "Konwersja modułu transmitancji do skali liniowej.")
    add_equation(doc, "Delta Z = |Z|lin · (10^(UD / 20) - 1)", "Błąd bezwzględny wynikający z niepewności amplitudy analizatora.")
    add_body(doc, "Dla każdego punktu wyznaczane są granice max(0, |Z|lin - Delta Z) oraz |Z|lin + Delta Z. Dla serii NEXT i FEXT program oblicza średnią, odchylenie standardowe z próby, błąd standardowy oraz przedział ufności średniej z rozkładu t-Studenta.")
    add_equation(doc, "CI95 = x̄ ± t(0,975; n-1) · s / √n", "Dla 11 punktów wartość krytyczna wynosi 2,228.")
    add_callout(doc, "Założenie", "Wartości początkowe UD wynoszą 0,2 dB dla 1-3 GHz i 0,3 dB dla 7-8 GHz, ale użytkownik może ustawić osobne U NEXT i U FEXT. W zastosowaniu metrologicznym należy je skonfrontować z poziomem sygnału i tabelą dokładności konkretnego analizatora.", "warning")

    add_heading(doc, "8. Algorytmy scenariusza sond pola bliskiego")
    add_equation(doc, "H[dBA/m] = P[dBm] - 30 + 10·log10(50) - K + Sp", "P - odczyt miernika mocy, K - wzmocnienie toru, Sp - poprawka sondy.")
    add_equation(doc, "H[A/m] = 10^(H[dBA/m] / 20)")
    add_equation(doc, "uH = √(uP² + uK² + uSp² + uRep²)", "Składnik powtarzalności uRep jest edytowany w pierwszym kroku i uwzględniany w budżecie.")
    add_equation(doc, "U95 = k · uH", "Domyślnie uP=0,066 dB, uK=0,2 dB, uSp=0,3 dB, k=2; U95≈0,733 dB.")
    add_body(doc, "Program wyznacza maksimum pola osobno dla linii 30 Ω, 50 Ω i 100 Ω. Szybkość zmian jest szacowana regresją liniową jako dB/100 MHz oraz dB/dekadę częstotliwości.")

    add_heading(doc, "9. Algorytmy scenariusza emisji promieniowanej EN55032")
    add_equation(doc, "AF = 20·log10(9,73 / (λ·√G)),    G = 1,64", "Poprawka antenowa dipola półfalowego dla polaryzacji poziomej.")
    add_equation(doc, "α = atan(hV / d)", "Kąt elewacji dla polaryzacji pionowej; domyślna odległość d = 3 m.")
    add_equation(doc, "E[dBµV/m] = MR[dBµV] + AF[dB/m] + IL[dB]", "Równanie przetwarzania zgodne z instrukcją ćwiczenia nr 3.")
    add_equation(doc, "UE = √(UMR² + UAF² + UIL²)", "Domyślnie UMR=0,2 dB, UAF=0,8 dB i UIL=0,2 dB.")
    add_body(doc, "Dla każdej częstotliwości program liczy pole dla obu polaryzacji, wybiera większą wartość, dodaje niepewność 95% i porównuje wynik z limitem EN55032 klasy B dla odległości 3 m. Limit wynosi 40 dBµV/m dla 30-230 MHz oraz 47 dBµV/m powyżej 230 MHz.")

    add_heading(doc, "10. Algorytmy scenariusza pomiarów propagacyjnych DVB-T")
    add_equation(doc, "U[µV] = 10^(-L[dB] / 20)", "Konwencja historyczna dla ujemnych odczytów zapisanych w sprawozdaniach z ćwiczenia nr 4.")
    add_equation(doc, "AF[1/m] = 10^(AF[dB/m] / 20)")
    add_equation(doc, "E[µV/m] = U[µV] · AF[1/m] · 10^(ac/20)", "ac oznacza tłumienie toru antenowego/kabla w dB.")
    add_equation(doc, "Eav ± T", "Wynik końcowy dla polaryzacji jest średnią przestrzenną z tolerancją wynikającą z rozrzutu punktów, niepewności odbiornika i AF.")
    add_body(doc, "Program obsługuje siatkę 16 punktów pomiarowych dla sygnału DVB-T około 522 MHz. Użytkownik wybiera konwencję danych wejściowych: zapis historyczny, dBµV albo dBm dla 50 Ω. AF jest interpolowany z tabeli UHALP 9108 A1 dla wybranego profilu lub może być podany ręcznie. Wyniki są liczone osobno dla obu polaryzacji i prezentowane również jako mapy cieplne 4 × 4.")

    add_heading(doc, "11. Walidacja i obsługa błędów")
    add_bullets(doc, [
        "NEXT i FEXT muszą mieścić się w zakresie od -200 dB do 0 dB.",
        "Moce w scenariuszu sond muszą mieścić się w zakresie od -150 dBm do 30 dBm.",
        "Wszystkie punkty wymagają kompletu wartości K i Sp.",
        "Krok przygotowania stanowiska wymaga zaznaczenia czterech pozycji listy kontrolnej.",
        "Scenariusz EN55032 wymaga kompletu MR, IL i wysokości anteny dla obu polaryzacji.",
        "Scenariusz propagacyjny wymaga 16 par odczytów dla polaryzacji poziomej i pionowej.",
        "Anulowanie okna zapisu nie zmienia stanu analizy.",
        "Import wymaga właściwej liczby punktów i poprawnych kolumn; błąd jest opisany przy tabeli.",
        "Błąd zapisu CSV lub DOCX jest prezentowany użytkownikowi w dolnym pasku okna.",
    ])

    add_heading(doc, "12. Import i eksport danych")
    add_body(doc, "MeasurementImportService odczytuje CSV/TXT z separatorem średnikowym, tabulatorem lub przecinkiem i obsługuje polski oraz niezmienny format liczb. Scenariusz nr 3 dodatkowo odczytuje pliki MATLAB z polami Pomiar_H, Pomiar_V, H2_H, H2_VH i IL. ReportGenerator zapisuje CSV UTF-8 z BOM. DocxReportGenerator tworzy raport Word z metadanymi, równaniami, tabelami wyników, podsumowaniem i miejscem na wnioski.")

    add_heading(doc, "13. Budowanie i publikowanie")
    add_code(doc, [
        "dotnet restore",
        "dotnet build CrosstalkAnalyzer.sln -c Release",
        "dotnet run --project CrosstalkAnalyzer.csproj",
        "dotnet publish -c Release -r win-x64 --self-contained true",
        "dotnet publish -c Release -r linux-x64 --self-contained true",
    ])
    add_callout(doc, "Zgodność", "Samowystarczalna publikacja nie wymaga instalacji .NET na komputerze docelowym, ale nadal podlega wymaganiom systemowym .NET 8 i Avalonia.", "info")

    add_heading(doc, "14. Testy")
    add_table(
        doc,
        ["Obszar", "Sprawdzenie", "Wynik"],
        [
            ("Konwersja dB", "-20 dB = 0,1; -40 dB = 0,01", "zaliczony"),
            ("Delta Z", "Wartość referencyjna dla 0,1 i UD=0,2 dB", "zaliczony"),
            ("Statystyka", "Średnia i s dla serii 1,2,3,4,5", "zaliczony"),
            ("Pole H", "Przykład 200 MHz, P=-34 dBm, K=21,25 dB, Sp=-31 dB", "zaliczony"),
            ("Niepewność", "uH≈0,366546 dB i U95≈0,733092 dB", "zaliczony"),
            ("EN55032", "AF 30 MHz, korekta pionowa i limit dla 200 MHz", "zaliczony"),
            ("Propagacja", "Konwersja -60,96 dB na 1116,863 µV i AF=23,04 dB/m", "zaliczony"),
            ("Import", "CSV czterech scenariuszy i struktura MATLAB", "zaliczony"),
            ("Eksport CSV", "Obecność wymaganych sekcji w czterech plikach", "zaliczony"),
            ("Eksport DOCX", "Dokument otwiera się i zawiera równania oraz tabele", "zaliczony"),
            ("Nawigacja", "Pełne przejście czterech kreatorów do kroku 4", "zaliczony"),
            ("UI headless", "820 × 600, brak kolizji przycisków i tekstu, widoczny skład wzorów", "zaliczony"),
            ("Kompilacja", "Release, 0 błędów i 0 ostrzeżeń", "zaliczony"),
        ],
        [1800, 5700, 1860],
    )
    add_body(doc, "Obliczenia i eksport są sprawdzane przez program kontrolny, a układ interfejsu przez Avalonia.Headless.XUnit i xUnit v3. Uruchomienie:")
    add_code(doc, [
        "dotnet run --project Tests/CrosstalkAnalyzer.CalculationChecks",
        "dotnet run --project Tests/CrosstalkAnalyzer.UiTests",
    ])

    add_heading(doc, "15. Ograniczenia i kierunki rozwoju")
    add_bullets(doc, [
        "Brak trwałego formatu sesji; po zamknięciu programu dane pozostają w wyeksportowanym CSV lub DOCX.",
        "Brak automatycznego odczytu K i Sp z pliku charakterystyki; wartości są zapisane jako edytowalne domyślne punkty.",
        "Ćwiczenia z ochrony środowiska oraz analizatora widma oczekują na instrukcje prowadzącego.",
        "Scenariusz propagacyjny wymaga zatwierdzenia względem oryginalnej instrukcji, której nie było w audytowanym katalogu.",
        "Brakuje wykładów cz02 i cz07; aplikacja nie rekonstruuje ich treści na podstawie numeracji.",
        "Możliwe rozszerzenie o zapis projektu, raport PDF i kolejne ćwiczenia po dostarczeniu źródeł.",
    ])

    add_heading(doc, "16. Struktura katalogów")
    add_code(doc, [
        "Models/      - rekordy wejściowe, wyniki i podsumowania",
        "Services/    - logika obliczeń, częstotliwości i eksport",
        "ViewModels/  - stan kreatorów, walidacja i nawigacja",
        "Views/       - widoki Avalonia XAML",
        "Controls/    - własne kontrolki wykresów",
        "Tests/       - kontrole obliczeń, eksportu, nawigacji i układu UI",
        "Documentation/COURSE_COVERAGE.md - rejestr pokrycia materiału",
        "Assets/      - zasoby aplikacji",
    ])

    add_heading(doc, "17. Materiały źródłowe")
    add_bullets(doc, [
        "Kod źródłowy projektu EMC Lab Assistant.",
        "R&S ZVL Vector Network Analyzer Data Sheet, wersja 12.00.",
        "R&S HZ-15/HZ-17 Probe Sets i R&S HZ-16 Preamplifier, Product Brochure.",
        "Instrukcja laboratoryjna „Pomiar sondami pola bliskiego”.",
        "„Obliczenia do pomiarów sondami pola bliskiego”.",
        "Wzorzec sprawozdania Lab_KEM_Pom_Emisji_Kabla_PCB.",
        "Instrukcja „Emisja promieniowana - poprawka antenowa, scenariusz pomiarowy normy EN55032”.",
        "Sprawozdanie i notatki z ćwiczenia „Pomiary propagacyjne (Nr 4)”.",
        "Rekomendacje ITU-R SM.1708-1, SM.1875-2, SM.378-7 oraz Spectrum Monitoring Handbook.",
        "Schwarzbeck UHALP 9108 A1, Correction for Short Measuring Distance.",
        "Materiały wykładowe i pytania kontrolne dostępne w audytowanym katalogu przedmiotu.",
    ])

    path = OUTPUT / "01_Dokumentacja_techniczna_EMC_Lab_Assistant.docx"
    doc.save(path)
    return path


def project_report():
    doc = Document()
    setup_document(doc, "Raport projektowy", "standard")
    add_cover(
        doc,
        "EMC Lab Assistant",
        "Wieloplatformowy kreator ćwiczeń laboratoryjnych z kompatybilności elektromagnetycznej",
        "Raport z projektu",
        [
            ("Przedmiot", COURSE),
            ("Autor", AUTHOR),
            ("Prowadzący", LECTURER),
            ("Wersja programu", VERSION),
            ("Technologia", "C# / .NET 8 / Avalonia UI"),
            ("Termin opracowania", DATE),
        ],
    )

    add_heading(doc, "Streszczenie")
    add_body(doc, "Celem projektu było wykonanie aplikacji w języku C#, która prowadzi studenta przez ćwiczenie laboratoryjne, prezentuje czytelnie złożone wzory, kontroluje kompletność danych i automatyzuje obliczenia. Zastosowano Avalonia UI, dzięki czemu rozwiązanie działa na systemach Windows i Linux. Wersja 2.0 obejmuje cztery scenariusze pomiarowe, moduł nauki, rejestr pokrycia przedmiotu, import CSV/MATLAB oraz eksport CSV/DOCX.")
    add_callout(doc, "Rezultat", "Powstał działający program z responsywnym interfejsem, czterema kreatorami, siedmioma blokami nauki, wykresami, mapami cieplnymi, obliczeniami niepewności i raportem DOCX.", "success")

    add_heading(doc, "1. Uzasadnienie projektu")
    add_body(doc, "Podczas laboratorium student wykonuje wiele powtarzalnych przeliczeń, korzysta z charakterystyk aparatury i musi zachować właściwą kolejność czynności. Błąd przy przepisywaniu danych, pominięcie poprawki lub zastosowanie niewłaściwej jednostki może prowadzić do niepoprawnego sprawozdania. Program pełni rolę asystenta: oddziela pomiar od obliczeń, pokazuje wzór na właściwym etapie oraz umożliwia wyeksportowanie tabeli wynikowej.")

    add_heading(doc, "2. Założenia i wymagania")
    add_bullets(doc, [
        "implementacja w języku C#;",
        "graficzny interfejs prowadzący użytkownika krok po kroku;",
        "działanie na Windows i Linux;",
        "prezentacja wzorów i automatyczne przeliczenia;",
        "walidacja danych wejściowych;",
        "wykresy i podsumowanie wyników;",
        "import danych z plików pomiarowych i eksport tabel oraz raportu;",
        "możliwość dołączania kolejnych scenariuszy laboratoryjnych.",
    ])

    add_heading(doc, "3. Wybór technologii")
    add_body(doc, "WPF i Windows Forms są silnie związane z systemem Windows. Avalonia UI wykorzystuje podobny model XAML i pozwala zachować język C# oraz wzorzec MVVM, a jednocześnie udostępnia backendy dla Windows i Linuksa. Projekt został ustawiony na .NET 8, czyli wspieraną wersję LTS, zamiast pierwotnego .NET 9.")
    add_table(
        doc,
        ["Kryterium", "Avalonia UI", "Wpływ na projekt"],
        [
            ("Wieloplatformowość", "Windows, Linux, macOS", "jeden kod interfejsu"),
            ("Język i XAML", "C# oraz deklaratywny XAML", "czytelny podział widok/logika"),
            ("MVVM", "pełne wsparcie wiązań", "łatwa walidacja i testowanie"),
            ("Dystrybucja", "publikacja self-contained", "brak wymogu instalacji runtime"),
        ],
        [2200, 2900, 4260],
    )

    add_heading(doc, "4. Organizacja rozwiązania")
    add_image(doc, ASSETS / "architektura.png", "Rysunek 1. Architektura projektu.")
    add_body(doc, "Rozdzielenie warstw ogranicza sprzężenie interfejsu z algorytmami. Wzory są testowane niezależnie od okna programu, a nowy scenariusz można dodać przez utworzenie modeli, usługi obliczeniowej, zestawu ViewModeli i odpowiadających im widoków.")

    add_heading(doc, "5. Zrealizowane scenariusze")
    add_image(doc, ASSETS / "przeplyw_scenariuszy.png", "Rysunek 2. Przebieg czterech kreatorów pomiarowych w wersji 2.0.")

    add_heading(doc, "5.1. Pomiar przeników między liniami mikropaskowymi", level=2)
    add_steps(doc, [
        "Wybór pasma 1–2 GHz, 2–3 GHz lub 7–8 GHz oraz wprowadzenie NEXT i FEXT.",
        "Konwersja z dB do skali liniowej.",
        "Obliczenie błędu wynikającego z niepewności analizatora.",
        "Obliczenie statystyk, 95% przedziału ufności, wykres i eksport.",
    ])
    add_equation(doc, "|Z|lin = 10^(|Z|dB/20)")
    add_equation(doc, "Delta Z = |Z|lin · (10^(UD/20) - 1)")

    add_heading(doc, "5.2. Sondy pola bliskiego w analizie emisji promieniowanej", level=2)
    add_steps(doc, [
        "Lista kontrolna stanowiska i zapis temperatury, wilgotności oraz ciśnienia.",
        "Wprowadzenie maksimum mocy dla linii 30 Ω, 50 Ω i 100 Ω w zakresie 100–1000 MHz.",
        "Zastosowanie wzmocnienia K i poprawki sondy Sp oraz obliczenie pola H.",
        "Porównanie charakterystyk, wyznaczenie maksimów, trendów i eksport.",
    ])
    add_equation(doc, "H[dBA/m] = P[dBm] - 30 + 10·log10(50) - K + Sp")
    add_equation(doc, "uH = √(uP² + uK² + uSp²),    U95 = 2·uH")

    add_heading(doc, "5.3. Emisja promieniowana - poprawka antenowa EN55032", level=2)
    add_steps(doc, [
        "Potwierdzenie odległości 3 m, dwóch polaryzacji i budżetu niepewności.",
        "Wprowadzenie MR, tłumienia kabla IL oraz wysokości anteny dla 30-1000 MHz.",
        "Wyznaczenie poprawki antenowej dipola i korekty dla polaryzacji pionowej.",
        "Porównanie maksymalnego pola E z limitem EN55032 klasy B i eksport.",
    ])
    add_equation(doc, "E[dBµV/m] = MR[dBµV] + AF[dB/m] + IL[dB]")
    add_equation(doc, "UE = √(UMR² + UAF² + UIL²)")

    add_heading(doc, "5.4. Pomiary propagacyjne DVB-T", level=2)
    add_steps(doc, [
        "Przygotowanie analizatora widma, anteny UHALP i pola pomiarowego 1 m × 1 m.",
        "Wpisanie poziomów dla polaryzacji poziomej i pionowej w 16 punktach.",
        "Przeliczenie ujemnych odczytów na U[µV] oraz pole E.",
        "Wyznaczenie średniej przestrzennej, tolerancji Eav ± T, wykresu i eksportu.",
    ])
    add_equation(doc, "U[µV] = 10^(-L[dB]/20)")
    add_equation(doc, "E[µV/m] = U[µV] · AF[1/m] · 10^(ac/20)")

    add_heading(doc, "6. Interfejs użytkownika")
    add_image(doc, ASSETS / "mapa_interfejsu.png", "Rysunek 3. Stałe elementy interfejsu programu.")
    add_body(doc, "Nagłówek informuje o aktywnym scenariuszu i kroku. Pasek postępu pokazuje pozycję w kreatorze. Środkowy panel ma pionowe przewijanie, formularze zawijają się przy mniejszej szerokości, a tabele przewijają się w obu kierunkach. Równania są składane przez CSharpMath zamiast wyświetlania surowego zapisu. Dolny pasek ma osobny wiersz statusu, dlatego przyciski nie zasłaniają treści.")

    add_heading(doc, "7. Weryfikacja")
    add_table(
        doc,
        ["Test", "Kryterium", "Rezultat"],
        [
            ("Kompilacja Release", "brak błędów i ostrzeżeń", "spełnione"),
            ("Uruchomienie okna", "proces nie kończy się przedwcześnie", "spełnione"),
            ("Wzory przeników", "wartości referencyjne", "spełnione"),
            ("Pole H", "-99,2603 dBA/m dla przykładu 200 MHz", "spełnione"),
            ("Niepewność sond", "U95≈0,7331 dB", "spełnione"),
            ("EN55032", "AF, korekta pionowa, limit i margines", "spełnione"),
            ("Propagacja", "konwersja odczytów, AF, Eav ± T", "spełnione"),
            ("Nawigacja", "pełne przejście czterech kreatorów", "spełnione"),
            ("Import", "CSV czterech scenariuszy i plik MATLAB", "spełnione"),
            ("Eksport", "wymagane sekcje CSV i raport DOCX", "spełnione"),
            ("UI 820 × 600", "brak kolizji przycisków/tekstów i widoczne wzory", "spełnione"),
        ],
        [2100, 4700, 2560],
    )

    add_heading(doc, "8. Wieloplatformowość")
    add_body(doc, "Projekt można opublikować jako samowystarczalną paczkę win-x64 lub linux-x64. Rozwiązanie nie wykorzystuje WPF, rejestru Windows ani innych mechanizmów specyficznych dla jednego systemu. Zapis plików odbywa się przez wieloplatformowy interfejs Avalonia StorageProvider.")
    add_callout(doc, "Granica zgodności", "Ze względu na wymagania .NET 8 aplikacja jest przeznaczona dla Windows 10/11. Windows 7 i Windows 8.1 nie są wspieranymi systemami docelowymi.", "warning")

    add_heading(doc, "9. Wartość dydaktyczna")
    add_bullets(doc, [
        "wymuszenie poprawnej kolejności działań laboratoryjnych;",
        "jawna prezentacja wzorów przed pokazaniem wyniku;",
        "oddzielenie danych surowych od danych skorygowanych;",
        "ułatwienie analizy niepewności i porównania serii;",
        "ograniczenie błędów rachunkowych i jednostkowych;",
        "możliwość wykorzystania eksportu jako podstawy tabel w sprawozdaniu.",
    ])

    add_heading(doc, "10. Ograniczenia")
    add_body(doc, "Program nie zastępuje oceny metrologicznej ani instrukcji obsługi aparatury. Domyślne współczynniki i niepewności należy weryfikować dla konkretnego stanowiska. Dwa ćwiczenia z sylabusa nie zostały zaimplementowane bez źródeł proceduralnych. Scenariusz propagacyjny pozostaje oznaczony jako wymagający zatwierdzenia względem oryginalnej instrukcji. Sesja nie jest zapisywana w formacie projektu.")

    add_heading(doc, "11. Możliwości dalszego rozwoju")
    add_bullets(doc, [
        "bezpośrednia komunikacja z analizatorem lub miernikiem;",
        "dodatkowe profile aparatury z plików kalibracyjnych;",
        "eksport raportu również do PDF;",
        "zapis i ponowne otwieranie sesji pomiarowej;",
        "wdrożenie dwóch brakujących laboratoriów po otrzymaniu instrukcji;",
        "uzupełnienie modułu nauki po otrzymaniu wykładów cz02 i cz07.",
    ])

    add_heading(doc, "12. Wnioski")
    add_body(doc, "Zrealizowano wymagany program w języku C# z wieloplatformowym interfejsem graficznym. Kreatory nie ograniczają się do kalkulatora: prowadzą przez przygotowanie pomiaru, kontrolują dane, prezentują podstawę matematyczną i tworzą wynik gotowy do dalszej analizy. Architektura MVVM i wydzielenie usług obliczeniowych umożliwiają rozszerzanie aplikacji bez przebudowy istniejących scenariuszy.")

    path = OUTPUT / "02_Raport_projektowy_dla_prowadzacego.docx"
    doc.save(path)
    return path


def user_manual():
    doc = Document()
    setup_document(doc, "Instrukcja użytkownika", "compact")
    add_cover(
        doc,
        "Instrukcja użytkowania",
        "Obsługa programu EMC Lab Assistant krok po kroku",
        "Podręcznik użytkownika",
        [
            ("Program", PROJECT),
            ("Wersja", VERSION),
            ("Odbiorca", "Student wykonujący ćwiczenie laboratoryjne"),
            ("Moduły", "4 kreatory; nauka; rejestr pokrycia materiału"),
            ("System", "Windows 10/11 lub Linux x64"),
            ("Autor", AUTHOR),
        ],
    )

    add_heading(doc, "1. Szybki start")
    add_steps(doc, [
        "Uruchom plik programu odpowiedni dla swojego systemu.",
        "Na ekranie startowym wybierz ćwiczenie.",
        "Wykonuj polecenia widoczne w kolejnych krokach.",
        "Po uzupełnieniu tabeli wybierz Dalej.",
        "Na ostatnim ekranie sprawdź wykres i użyj Eksportuj CSV lub Eksportuj DOCX.",
    ])
    add_callout(doc, "Ważne", "Przycisk Dalej jest nieaktywny, dopóki wszystkie wymagane pola nie są poprawnie uzupełnione.", "info")

    add_heading(doc, "2. Wymagania i uruchomienie")
    add_table(
        doc,
        ["Wariant", "Sposób uruchomienia"],
        [
            ("Paczka samowystarczalna", "Uruchom CrosstalkAnalyzer.exe w Windows albo plik CrosstalkAnalyzer w Linux."),
            ("Kod źródłowy", "W katalogu projektu wykonaj polecenie: dotnet run."),
        ],
        [2700, 6660],
    )
    add_body(doc, "Program zapisuje wyłącznie pliki wybrane przez użytkownika. Nie wymaga połączenia z Internetem i nie wysyła danych pomiarowych.")

    add_heading(doc, "3. Układ okna")
    add_image(doc, ASSETS / "mapa_interfejsu.png", "Rysunek 1. Rozmieszczenie najważniejszych elementów.")
    add_table(
        doc,
        ["Element", "Znaczenie"],
        [
            ("Nagłówek", "Nazwa kroku, scenariusza oraz krótka instrukcja."),
            ("Pasek postępu", "Pozycja od 1 do 4 w bieżącym scenariuszu."),
            ("Obszar roboczy", "Lista kontrolna, tabela danych, wzór, wyniki lub wykres."),
            ("Wstecz", "Powrót do poprzedniego kroku bez usuwania danych."),
            ("Zmień scenariusz", "Powrót do listy ćwiczeń."),
            ("Dalej", "Przeliczenie danych i przejście do kolejnego kroku."),
            ("Eksportuj CSV", "Zapis danych do dalszej analizy w arkuszu."),
            ("Eksportuj DOCX", "Zapis sformatowanego raportu z równaniami, tabelami i miejscem na wnioski."),
        ],
        [2300, 7060],
    )

    add_heading(doc, "4. Wybór scenariusza")
    add_body(doc, "Ekran startowy zawiera cztery karty ćwiczeń oraz dwie karty pomocnicze: Nauka i Pokrycie materiału. Wybierz kreator odpowiadający wykonywanemu ćwiczeniu albo moduł pomocniczy. Zmiana modułu jest możliwa w każdej chwili przyciskiem Zmień scenariusz.")
    add_image(doc, ASSETS / "przeplyw_scenariuszy.png", "Rysunek 2. Kolejność kroków w czterech scenariuszach.")

    add_heading(doc, "5. Scenariusz: pomiar przeników")
    add_heading(doc, "Krok 1. Dane wejściowe", level=2)
    add_steps(doc, [
        "Wybierz pasmo 1–2 GHz, 2–3 GHz albo 7–8 GHz.",
        "Program utworzy 11 punktów częstotliwości.",
        "Wpisz przenik bliski NEXT i przenik daleki FEXT w dB.",
        "Używaj wartości od -200 dB do 0 dB.",
        "Opcjonalnie wybierz Wstaw dane przykładowe, aby zaprezentować działanie programu.",
    ])
    add_callout(doc, "Format liczb", "Można wpisywać liczby dziesiętne zgodnie z ustawieniami regionalnymi systemu. Wartości przeników są zwykle ujemne.", "info")

    add_heading(doc, "Krok 2. Konwersja", level=2)
    add_body(doc, "Program pokazuje dane wejściowe i ich odpowiedniki w skali liniowej. Ekran jest tylko do odczytu.")
    add_equation(doc, "|Z|lin = 10^(|Z|dB/20)")

    add_heading(doc, "Krok 3. Błąd analizatora", level=2)
    add_body(doc, "Sprawdź osobno wartości U NEXT i U FEXT oraz przedziały każdego punktu. Wartości domyślne wynoszą 0,2 dB dla pasm do 3 GHz i 0,3 dB dla pasma 7-8 GHz, ale można je edytować zgodnie ze specyfikacją analizatora.")
    add_equation(doc, "Delta Z = |Z|lin · (10^(UD/20) - 1)")

    add_heading(doc, "Krok 4. Podsumowanie", level=2)
    add_body(doc, "Ekran zawiera statystyki NEXT/FEXT oraz wykres przebiegów. Przedział ufności dotyczy średniej z punktów w wybranym paśmie.")
    add_bullets(doc, [
        "Nowe badanie - czyści dane i wraca do pierwszego kroku.",
        "Eksportuj CSV - zapisuje punkty, błędy, granice oraz statystyki.",
        "Eksportuj DOCX - tworzy raport gotowy do uzupełnienia wniosków.",
    ])

    add_heading(doc, "6. Scenariusz: sondy pola bliskiego")
    add_heading(doc, "Krok 1. Przygotowanie stanowiska", level=2)
    add_steps(doc, [
        "Połącz sondę R&S H 400-1 przez wzmacniacz z miernikiem mocy.",
        "Ustaw generator na sygnał niemodulowany i moc 10 dBm.",
        "Ustaw w mierniku częstotliwość zgodną z generatorem.",
        "Zapisz temperaturę, wilgotność i ciśnienie.",
        "Zaznacz wszystkie pozycje listy kontrolnej.",
    ])
    add_callout(doc, "Technika pomiaru", "Pętlę sondy przystaw ściśle do kabla. Przesuwaj ją wzdłuż całego kabla i obracaj tak, aby znaleźć maksimum globalne.", "warning")

    add_heading(doc, "Krok 2. Pomiary", level=2)
    add_body(doc, "Dla częstotliwości 100–1000 MHz wpisz moc P w dBm dla linii 30 Ω, 50 Ω i 100 Ω. Kolumny K i Sp są wypełnione wartościami odczytanymi z materiałów ćwiczenia, ale pozostają edytowalne.")
    add_table(
        doc,
        ["Kolumna", "Co wpisać"],
        [
            ("P - 30 Ω", "Maksimum wskazania dla płytki 30 Ω."),
            ("P - 50 Ω", "Maksimum wskazania dla płytki 50 Ω."),
            ("P - 100 Ω", "Maksimum wskazania dla płytki 100 Ω."),
            ("K", "Wzmocnienie toru pomiarowego w dB."),
            ("Sp", "Poprawka sondy pola bliskiego w dB."),
        ],
        [2600, 6760],
    )
    add_callout(doc, "Dane przykładowe", "Przycisk służy do demonstracji i testowania programu. Przed przygotowaniem sprawozdania zastąp te wartości własnymi pomiarami.", "warning")

    add_heading(doc, "Krok 3. Pole H i niepewność", level=2)
    add_equation(doc, "H[dBA/m] = P[dBm] - 30 + 10·log10(50) - K + Sp")
    add_equation(doc, "H[A/m] = 10^(H[dBA/m]/20)")
    add_body(doc, "Program pokazuje wyniki dla każdej impedancji oraz 95% przedziały. Budżet obejmuje uP, uK, uSp i edytowalny składnik powtarzalności uRep.")

    add_heading(doc, "Krok 4. Porównanie", level=2)
    add_body(doc, "Tabela podsumowuje największą wartość H, odpowiadającą częstotliwość oraz oszacowaną szybkość zmian. Wykres umożliwia porównanie linii 30 Ω, 50 Ω i 100 Ω. Pola Obserwacje i Wnioski służą do zapisania analizy dodatkowego nagrania.")
    add_callout(doc, "Interpretacja trendu", "Wartość dodatnia w kolumnie dB/100 MHz oznacza, że w całym badanym zakresie poziom H ma tendencję rosnącą. Lokalne maksima i spadki nadal mogą występować.", "info")

    add_heading(doc, "7. Scenariusz: emisja promieniowana EN55032")
    add_heading(doc, "Krok 1. Założenia", level=2)
    add_steps(doc, [
        "Potwierdź odległość pomiarową 3 m i zakres 30-1000 MHz.",
        "Zaznacz, że polaryzacja pozioma korzysta ze standardowej poprawki AF.",
        "Zaznacz, że polaryzacja pionowa wymaga korekty geometrycznej.",
        "Sprawdź domyślne niepewności: MR 0,2 dB, AF 0,8 dB, IL 0,2 dB.",
    ])
    add_callout(doc, "Korekta pionowa", "W tym scenariuszu kierunek odbioru dla polaryzacji pionowej nie pokrywa się z maksimum charakterystyki dipola, dlatego program liczy kąt elewacji i dodatkową korektę AF.", "info")

    add_heading(doc, "Krok 2. Dane surowe", level=2)
    add_body(doc, "W tabeli wpisz tłumienie kabla IL, wskazanie analizatora MR oraz wysokość anteny, przy której odczyt był maksymalny. Dane wpisuje się osobno dla polaryzacji poziomej i pionowej. Przycisk Importuj MAT/CSV odczytuje także strukturę Data z oryginalnych plików MATLAB.")
    add_table(
        doc,
        ["Kolumna", "Znaczenie"],
        [
            ("IL", "Tłumienie kabla lub toru pomiarowego w dB."),
            ("MR H / MR V", "Odczyt analizatora w dBµV dla polaryzacji poziomej lub pionowej."),
            ("h H / h V", "Wysokość anteny odbiorczej w metrach."),
        ],
        [2600, 6760],
    )

    add_heading(doc, "Krok 3. Poprawka antenowa i pole E", level=2)
    add_equation(doc, "AF = 20·log10(9,73 / (λ·√G))")
    add_equation(doc, "E[dBµV/m] = MR[dBµV] + AF[dB/m] + IL[dB]")
    add_body(doc, "Ekran pokazuje AF dla polaryzacji poziomej, kąt α, korektę pionową, AF skorygowane oraz pola E dla obu polaryzacji.")

    add_heading(doc, "Krok 4. Limit i margines", level=2)
    add_body(doc, "Program wybiera większe pole z obu polaryzacji, tworzy 95% przedział ufności i porównuje górną granicę z limitem EN55032 klasy B. Dodatni margines oznacza przekroczenie, a wartość ujemna oznacza zapas.")
    add_callout(doc, "Limity", "Dla odległości 3 m program przyjmuje 40 dBµV/m w zakresie 30-230 MHz oraz 47 dBµV/m powyżej 230 MHz.", "warning")

    add_heading(doc, "8. Scenariusz: pomiary propagacyjne DVB-T")
    add_heading(doc, "Krok 1. Stanowisko", level=2)
    add_steps(doc, [
        "Ustaw analizator widma w zakresie 200 MHz - 1 GHz.",
        "Zidentyfikuj sygnał DVB-T w okolicy 522 MHz.",
        "Przyjmij pole pomiarowe 1 m × 1 m i 16 punktów.",
        "Wybierz profil AF UHALP 9108 albo tryb ręczny oraz właściwą konwencję odczytu: historyczną, dBµV lub dBm/50 Ω.",
    ])

    add_heading(doc, "Krok 2. Tabela 16 punktów", level=2)
    add_body(doc, "Wpisz poziom z analizatora dla polaryzacji poziomej i pionowej. W każdym punkcie należy obrócić antenę w celu znalezienia maksimum.")
    add_callout(doc, "Dane przykładowe", "Domyślne dane z raportu dotyczą sygnału DVB-T około 522 MHz. Do własnego sprawozdania należy wpisać własne wyniki pomiarów.", "warning")

    add_heading(doc, "Krok 3. Przeliczenie na pole E", level=2)
    add_equation(doc, "U[µV] = 10^(-L[dB] / 20)")
    add_equation(doc, "E[µV/m] = U[µV] · AF[1/m] · 10^(ac/20)")
    add_body(doc, "Program liczy wartości U i E osobno dla obu polaryzacji oraz wskazuje silniejszą polaryzację w każdym punkcie.")

    add_heading(doc, "Krok 4. Średnia i tolerancja", level=2)
    add_body(doc, "Ostatni ekran pokazuje wynik końcowy w postaci Eav ± T, punkt maksimum, wykres polaryzacji poziomej i pionowej oraz dwie mapy cieplne siatki 4 × 4.")
    add_callout(doc, "Interpretacja", "Jeżeli polaryzacja pozioma ma wyższy średni poziom pola, jest to zgodne z oczekiwaniem dla sygnału DVB-T nadawanego w polaryzacji poziomej.", "info")

    add_heading(doc, "9. Import i eksport")
    add_steps(doc, [
        "W kroku danych wybierz Importuj CSV lub Importuj MAT/CSV i wskaż plik.",
        "Sprawdź komunikat importu oraz kompletność tabeli.",
        "Przejdź do czwartego kroku.",
        "Wybierz Eksportuj CSV albo Eksportuj DOCX.",
        "Wskaż katalog i nazwę pliku.",
        "CSV otwórz w arkuszu, a DOCX w Wordzie albo LibreOffice Writer.",
    ])
    add_body(doc, "Każdy scenariusz eksportuje dane surowe, wartości pośrednie oraz podsumowanie. Raport DOCX dodaje zastosowane równania, metadane ćwiczenia i miejsce na końcową interpretację.")

    add_heading(doc, "10. Nauka i pokrycie materiału")
    add_body(doc, "Karta Nauka zawiera siedem bloków tematycznych, pytania kontrolne oraz kalkulatory długości fali, prądu pojemnościowego i napięcia indukowanego. Karta Pokrycie materiału pokazuje ćwiczenia i wykłady, których nie można uznać za wdrożone bez dodatkowych źródeł.")
    add_callout(doc, "Uwaga", "Pozycja Oczekuje na instrukcję nie jest awarią programu. Chroni przed użyciem procedury lub wzoru odtworzonego bez podstawy źródłowej.", "warning")

    add_heading(doc, "11. Najczęstsze problemy")
    add_table(
        doc,
        ["Objaw", "Przyczyna", "Rozwiązanie"],
        [
            ("Dalej jest nieaktywne", "Brakuje danych lub pozycja listy kontrolnej nie jest zaznaczona.", "Sprawdź komunikat pod tabelą i wszystkie pola."),
            ("Wartość nie zostaje przyjęta", "Nieprawidłowy format lub zakres liczby.", "Usuń jednostkę z pola; wpisz wyłącznie liczbę."),
            ("Wynik wygląda nietypowo", "Nieprawidłowy znak K/Sp, AF/IL albo jednostka wejścia.", "Sprawdź, czy P jest w dBm, MR w dBµV, a odczyty propagacyjne używają właściwej konwencji."),
            ("CSV jest w jednej kolumnie", "Arkusz nie rozpoznał separatora.", "Zaimportuj plik z separatorem średnikowym."),
            ("Import został odrzucony", "Brakuje kolumn, punktów lub liczba ma niepoprawny format.", "Sprawdź nagłówek i wymaganą liczbę wierszy dla scenariusza."),
            ("Nie można zapisać pliku", "Brak uprawnień lub plik jest otwarty.", "Wybierz inny katalog/nazwę i zamknij plik w arkuszu."),
        ],
        [2100, 3400, 3860],
    )

    add_heading(doc, "12. Dobre praktyki")
    add_bullets(doc, [
        "Zapisuj CSV i DOCX natychmiast po zakończeniu serii pomiarowej.",
        "Nie traktuj danych przykładowych jako wyników laboratoryjnych.",
        "Przed eksportem sprawdź znaki i jednostki wszystkich współczynników.",
        "Zachowuj stałą orientację sondy albo anteny podczas porównywania serii.",
        "W sprawozdaniu podaj zastosowane wzory i źródła niepewności.",
        "Porównuj wyniki z instrukcją oraz kartą katalogową użytej aparatury.",
    ])

    add_heading(doc, "13. Skrócona lista kontrolna przed oddaniem sprawozdania")
    add_bullets(doc, [
        "Użyto własnych danych pomiarowych.",
        "Wartości K, Sp, UD, AF, IL i niepewności zostały zweryfikowane.",
        "Tabela zawiera jednostki.",
        "Wykres ma podpisane osie i legendę.",
        "Podano niepewność i sposób jej obliczenia.",
        "Plik CSV albo raport DOCX został zachowany wraz ze sprawozdaniem.",
    ])

    path = OUTPUT / "03_Instrukcja_uzytkowania_EMC_Lab_Assistant.docx"
    doc.save(path)
    return path


def main():
    create_diagrams()
    paths = [technical_document(), project_report(), user_manual()]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
